import itertools
from datetime import date
from io import BytesIO

from django.contrib.auth.decorators import login_required
from django.db.models import QuerySet, Avg
from django.db.models.functions import Round
from django.http import HttpResponseRedirect, HttpResponse
from django.shortcuts import render

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from children.models import ChildrenGroup, Children
from djangoProject.forms import LessonForm, ResultForm, SurveyForm, ScopeForm
from schedule.models import Lesson, Survey, ChildrenScopeResult, Scope


@login_required
def group_schedule_detail(request, pk):
    """Вывод информации о расписании группы"""
    group = ChildrenGroup.objects.get(id=pk)
    list_lesson = Lesson.objects.filter(Group=pk)
    access = request.user.groups.filter(name="Старший воспитатель").exists() or \
             request.user.groups.filter(name="Заведующий").exists()
    if request.method == "GET":
        if 'month' in request.GET:
            month_request = request.GET.get("month")
            list_lesson = list_lesson.filter(Date__month=month_request)

    if request.method == "POST":
        if 'add_form' in request.POST:
            form = LessonForm(request.POST)
            if form.is_valid():
                instance = form.save(commit=False)
                instance.Group = group
                if not access:
                    instance.Score = 0
                instance.save()

    context = {
        'list_lesson': list_lesson,
        'form': LessonForm(),
        'group': group,
        'access': access
    }
    return render(request, 'schedule/schedule_detail.html', context)


@login_required
def group_schedule(request):
    """Вывод учебных групп"""
    access = request.user.groups.filter(name="Старший воспитатель").exists() or \
             request.user.groups.filter(name="Заведующий").exists()

    if access: group_list = ChildrenGroup.objects.all()
    else: group_list = ChildrenGroup.objects.filter(Mentor=request.user.id)

    context = {
        'group_list': group_list,
    }
    return render(request, 'schedule/schedule.html', context)

@login_required
def survey(request):
    """Вывод учебных групп"""
    access = request.user.groups.filter(name="Старший воспитатель").exists() or \
             request.user.groups.filter(name="Заведующий").exists()

    if access: group_list = ChildrenGroup.objects.all()
    else: group_list = ChildrenGroup.objects.filter(Mentor=request.user.id)

    context = {
        'list_group': group_list,

    }
    return render(request, 'schedule/survey.html', context)

@login_required
def group_survey(request, pk):
    group = ChildrenGroup.objects.get(id=pk)
    survey_list = Survey.objects.filter(Group=pk)
    if request.method == "POST":
        if 'add_form' in request.POST:
            form = SurveyForm(request.POST)
            if form.is_valid():
                instance = form.save(commit=False)
                instance.Group = group
                instance.save()
    context = {
        'survey_list': survey_list,
        'group': group,
        'form': SurveyForm
    }
    return render(request, 'schedule/group_survey.html', context)

@login_required
def survey_delete(request, pk):
    get_lesson = Survey.objects.get(pk=pk)
    get_lesson.delete()
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

@login_required
def survey_detail(request, pk):
    SurveyInfo = Survey.objects.get(id=pk)
    SurveyScope = Scope.objects.filter(Survey=pk)
    ChildrenList = Children.objects.filter(Group=pk)
    ChildrenResultList = ChildrenScopeResult.objects.filter(Survey_id = pk)

    avg = ChildrenScopeResult.objects.filter(Survey_id = pk).values('Scope').annotate(avg = Avg('Result'))
    childAvg = ChildrenScopeResult.objects.filter(Survey_id=pk).values('Children').annotate(avg=Round(Avg('Result'),1))
    if request.method == "POST":
        if 'add_form' in request.POST:
            form = ResultForm(request.POST)
            if form.is_valid():
                instance = form.save(commit=False)
                instance.Group = ChildrenGroup.objects.get(id=pk)
                instance.Survey = SurveyInfo
                instance.save()

        if 'add_form1' in request.POST:
            form = ScopeForm(request.POST)
            if form.is_valid():
                instance = form.save(commit=False)
                instance.Survey = SurveyInfo
                instance.save()

    context = {
        'SurveyInfo': SurveyInfo,
        'SurveyScope': SurveyScope,
        'ChildrenList': ChildrenList,
        'ChildrenResultList': ChildrenResultList,
        'avg': avg,
        'childAvg': childAvg,
        'form': ResultForm,
        'Scopeform': ScopeForm,
    }
    return render(request, 'schedule/survey_detail.html', context)


@login_required
def survey_detail_info(request, pk):
    SurveyInfo = Survey.objects.get(id=pk)
    context = {
        'SurveyInfo': SurveyInfo,
    }
    return render(request, 'schedule/survey_detail_info.html', context)

@login_required
def lesson_detail(request, pk):
    LessonInfo = Lesson.objects.get(id=pk)
    context = {
        'LessonInfo': LessonInfo,
    }
    return render(request, 'schedule/lesson.html', context)

@login_required
def lesson_delete(request, pk):
    get_lesson = Lesson.objects.get(pk=pk)
    get_lesson.delete()
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

@login_required
def lesson_update(request, pk):
    get_lesson = Lesson.objects.get(pk=pk)
    get_group = ChildrenGroup.objects.get(pk=get_lesson.Group.pk)
    access = request.user.groups.filter(name="Старший воспитатель").exists() or \
             request.user.groups.filter(name="Заведующий").exists()
    if request.method == "POST":
        form = LessonForm(request.POST, instance=get_lesson)
        if form.is_valid():
            instance = form.save(commit=False)
            if not access:
               instance.Score = 0
            instance.Group = get_group
            instance.save()
    context = {
        'get_lesson': get_lesson,
        'update': True,
        'form': LessonForm(instance=get_lesson),
        'access': access,
    }
    return render(request, 'schedule/lesson_update.html', context)


def group_schedule_docx(request, pk):
    list_lesson = Lesson.objects.all().filter(Group=pk).query
    list_lesson.group_by = ['Activities']
    results = list(QuerySet(query=list_lesson, model=Lesson))
    document = Document()
    docx_title="Перспективно-тематическое планирование образовательной деятельности"

    # ---- Cover Letter ----
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    tittle = document.add_paragraph()
    tittle.add_run("Перспективно-тематическое планирование образовательной деятельности").bold = True
    tittle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prev = 0
    for i in results:

        if i.Activities != prev:
            header = document.add_paragraph()
            header.add_run(str(i.Activities)).bold = True
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

            table = document.add_table(rows=1, cols=6, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Тема'
            hdr_cells[2].text = 'Задачи'
            hdr_cells[3].text = 'Организация предметно–пространственной развивающей среды'
            hdr_cells[4].text = 'Используемая литература'
            hdr_cells[5].text = 'Виды деятельности и культурные практики'

            row_cells = table.add_row().cells
            row_cells[0].text = i.Date.strftime('%d.%m.%Y')
            row_cells[1].text = str(i.Theme)
            row_cells[2].text = str(i.Task)
            row_cells[3].text = str(i.Material)
            row_cells[4].text = str(i.Literature)
            row_cells[5].text = str(i.Activities)
            prev = i.Activities
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = i.Date.strftime('%d.%m.%Y')
            row_cells[1].text = str(i.Theme)
            row_cells[2].text = str(i.Task)
            row_cells[3].text = str(i.Material)
            row_cells[4].text = str(i.Literature)
            row_cells[5].text = str(i.Activities)

    # Prepare document for download
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

